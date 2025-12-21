from docx import Document

# Create the Word document
doc = Document()
doc.add_heading('Interface Control Document (ICD)', 0)
doc.add_heading('Connector: MIL-DTL-D38999/20WC35PN (22-Pin)', level=1)

doc.add_paragraph(
    "This Interface Control Document (ICD) defines the electrical and mechanical interface between "
    "the Master (control card) and the Slave (subcontractor's device). Communication is based on "
    "SPI over RS-422 differential signaling with discrete and debug interfaces. The physical connector "
    "used is the MIL-DTL-D38999/20WC35PN, a 22-pin circular connector."
)

doc.add_heading('1. Connector Pin Assignment', level=2)

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Pin #'
hdr_cells[1].text = 'Signal Name'
hdr_cells[2].text = 'Direction'
hdr_cells[3].text = 'Type'
hdr_cells[4].text = 'Description'

pins = [
    (1, 'SPI_CLK+', 'Master → Slave', 'RS-422 differential', 'SPI Clock (+)'),
    (2, 'SPI_CLK−', 'Master → Slave', 'RS-422 differential', 'SPI Clock (−)'),
    (3, 'SPI_MOSI+', 'Master → Slave', 'RS-422 differential', 'SPI Master Out, Slave In (+)'),
    (4, 'SPI_MOSI−', 'Master → Slave', 'RS-422 differential', 'SPI Master Out, Slave In (−)'),
    (5, 'SPI_MISO+', 'Slave → Master', 'RS-422 differential', 'SPI Slave Out, Master In (+)'),
    (6, 'SPI_MISO−', 'Slave → Master', 'RS-422 differential', 'SPI Slave Out, Master In (−)'),
    (7, 'SPI_CSN+', 'Master → Slave', 'RS-422 differential', 'Chip Select Not (+), active low'),
    (8, 'SPI_CSN−', 'Master → Slave', 'RS-422 differential', 'Chip Select Not (−)'),
    (9, 'RESET_N+', 'Master → Slave', 'RS-422 differential', 'Reset (active low) (+)'),
    (10, 'RESET_N−', 'Master → Slave', 'RS-422 differential', 'Reset (active low) (−)'),
    (11, 'RF_RDY+', 'Slave → Master', 'RS-422 differential', 'RF Ready signal (+)'),
    (12, 'RF_RDY−', 'Slave → Master', 'RS-422 differential', 'RF Ready signal (−)'),
    (13, 'GPIO_RS422_OUT+', 'Master → Slave', 'RS-422 differential', 'Reserved GPIO output (+)'),
    (14, 'GPIO_RS422_OUT−', 'Master → Slave', 'RS-422 differential', 'Reserved GPIO output (−)'),
    (15, 'GPIO_RS422_IO+', 'Bidirectional', 'RS-422 differential', 'Reserved bidirectional GPIO (+)'),
    (16, 'GPIO_RS422_IO−', 'Bidirectional', 'RS-422 differential', 'Reserved bidirectional GPIO (−)'),
    (17, 'RS232_TXD', 'Slave → Master', 'RS-232', 'Serial debug TX from Slave'),
    (18, 'RS232_RXD', 'Master → Slave', 'RS-232', 'Serial debug RX to Slave'),
    (19, 'DEBUG_EN', 'Master → Slave', 'LVTTL / Discrete', 'Enable RS-232 debug on Slave'),
    (20, 'GND_RS422', '—', 'Ground', 'Return for all RS-422 signals'),
    (21, 'GND_RS232', '—', 'Ground', 'Return for RS-232 and DEBUG_EN'),
    (22, 'SHIELD_GND', '—', 'Shield/Chassis GND', 'Cable shield connection'),
]

for pin in pins:
    row_cells = table.add_row().cells
    for i, value in enumerate(pin):
        row_cells[i].text = str(value)

doc.add_heading('2. Connector Diagram', level=2)
doc.add_paragraph(
    "The diagram below illustrates the pin numbering for the MIL-DTL-D38999/20WC35PN connector "
    "(front face view, looking into receptacle):"
)
doc.add_paragraph("[Connector pin diagram image here — insert manually or use project CAD data]")

doc.save("ICD_D38999_22Pin_SPI_RS422_RS232.docx")
